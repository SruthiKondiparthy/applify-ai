def calculate_match_percentage(resume_text: str, job_text: str) -> float:
    """
    Calculates a simple match percentage between resume and job description using token overlap.
    Returns a value between 0 and 100.
    """
    import re
    resume_tokens = set(re.findall(r'\w+', resume_text.lower()))
    job_tokens = set(re.findall(r'\w+', job_text.lower()))
    if not job_tokens:
        return 0.0
    overlap = resume_tokens & job_tokens
    return round(100 * len(overlap) / len(job_tokens), 2)
# api/utils.py
from io import BytesIO
from typing import Tuple
from datetime import datetime
from fpdf import FPDF
from docx import Document

def create_pdf_from_text(title: str, body: str) -> bytes:
    """
    Simple PDF generator using FPDF. For production use, enhance fonts and layout.
    Returns raw PDF bytes.
    """
    pdf = FPDF(format='A4')
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 6, title)
    pdf.ln(4)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, body)
    output = BytesIO()
    pdf.output(output)
    return output.getvalue()

def create_docx_from_text(title: str, body: str) -> bytes:
    """
    Create a simple DOCX file.
    """
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body.splitlines():
        if line.strip() == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
